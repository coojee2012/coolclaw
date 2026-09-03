"""DOCX template skill — generate or fill Word documents using python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def run(
    output: str,
    template: str = "",
    data: dict | None = None,
    content: str = "",
) -> dict:
    """生成或填充 Word 文档。

    Args:
        output: 输出文件路径
        template: 模板文件路径（留空则创建新文档）
        data: 模板变量 {"key": "value"}
        content: 要写入的文本内容

    Returns:
        生成结果
    """
    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if template and Path(template).exists():
        # Fill template
        result = _fill_template(template, data or {}, str(output_path))
    elif content:
        # Create from content
        result = _create_from_content(content, str(output_path))
    else:
        # Create empty document
        result = _create_empty(str(output_path))

    return result


def _fill_template(template_path: str, data: dict, output_path: str) -> dict:
    """Fill a Word template with data.

    Supports {{variable}} placeholders in paragraphs and tables.
    """
    doc = Document(template_path)
    filled_count = 0

    # Process paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            new_text, count = _replace_placeholders(text, data)
            if count > 0:
                run.text = new_text
                filled_count += count

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        new_text, count = _replace_placeholders(text, data)
                        if count > 0:
                            run.text = new_text
                            filled_count += count

    doc.save(output_path)
    return {
        "status": "ok",
        "output": output_path,
        "template": template_path,
        "filled_placeholders": filled_count,
        "variables_used": list(data.keys()),
    }


def _replace_placeholders(text: str, data: dict) -> tuple[str, int]:
    """Replace {{key}} placeholders in text. Returns (new_text, count)."""
    count = 0

    def replacer(match: re.Match) -> str:
        nonlocal count
        key = match.group(1).strip()
        if key in data:
            count += 1
            return str(data[key])
        return match.group(0)  # keep original if no match

    new_text = re.sub(r"\{\{(.+?)\}\}", replacer, text)
    return new_text, count


def _create_from_content(content: str, output_path: str) -> dict:
    """Create a new Word document from text content."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Add content as paragraphs
    lines = content.split("\n")
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # empty line

    doc.save(output_path)
    return {
        "status": "ok",
        "output": output_path,
        "paragraphs": len([l for l in lines if l.strip()]),
        "total_lines": len(lines),
    }


def _create_empty(output_path: str) -> dict:
    """Create an empty Word document."""
    doc = Document()
    doc.save(output_path)
    return {
        "status": "ok",
        "output": output_path,
        "message": "Empty document created",
    }
